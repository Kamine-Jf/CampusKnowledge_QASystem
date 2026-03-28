"""Document ingestion service for Excel, PDF and Word uploads."""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import List

from docx import Document
from fastapi import UploadFile

from src.config.settings import PdfConfig, PROJECT_ROOT
from src.database.db_operate import add_pdf_info, list_pdf_info, del_pdf_info


STRUCTURED_DIR = PROJECT_ROOT / "data" / "structured_data"
UNSTRUCTURED_DIR = PROJECT_ROOT / "data" / "unstructured_data"

ALLOWED_EXCEL = {".xlsx", ".xls"}
ALLOWED_UNSTRUCTURED = {".pdf", ".docx", ".doc"}
ALLOWED_ALL = ALLOWED_EXCEL | ALLOWED_UNSTRUCTURED


def _validate_file_name(file_name: str | None) -> str:
    """Validate uploaded file name and return safe name."""
    normalized_name = (file_name or "").strip()
    if not normalized_name:
        raise ValueError("file name is required")

    suffix = Path(normalized_name).suffix.lower()
    if suffix not in ALLOWED_ALL:
        raise ValueError(f"unsupported file type: {suffix}, supported: {', '.join(sorted(ALLOWED_ALL))}")

    return Path(normalized_name).name


def _get_target_dir(file_name: str) -> Path:
    """Route file to correct directory based on extension."""
    suffix = Path(file_name).suffix.lower()
    if suffix in ALLOWED_EXCEL:
        return STRUCTURED_DIR
    return UNSTRUCTURED_DIR


def _persist_file_bytes(file_name: str, content: bytes) -> Path:
    """Save uploaded file to the appropriate directory."""
    target_dir = _get_target_dir(file_name)
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / file_name
    target_path.write_bytes(content)
    return target_path


def _parse_docx_to_chunks(local_path: Path) -> List[str]:
    """Parse docx text and convert to semantic chunks."""
    from src.vector_db.pdf2vector import _chunk_page_text, preprocess_text

    doc = Document(str(local_path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text and paragraph.text.strip()]
    if not paragraphs:
        return []

    joined_text = preprocess_text("\n\n".join(paragraphs))
    chunks, _, _ = _chunk_page_text(
        joined_text,
        PdfConfig.MIN_CHARS_PER_CHUNK,
        PdfConfig.MAX_CHARS_PER_CHUNK,
    )
    return chunks


def _parse_doc_to_chunks(local_path: Path) -> List[str]:
    """Parse legacy .doc (Word 97-2003) binary format to text chunks.

    Try chain:
    1. python-docx  — works when the .doc is actually OOXML-wrapped
    2. win32com     — Windows + Microsoft Office COM automation
    3. Raise ValueError with a helpful message
    """
    from src.vector_db.pdf2vector import _chunk_page_text, preprocess_text

    # --- attempt 1: python-docx (handles OOXML files renamed to .doc) ---
    try:
        doc = Document(str(local_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        if paragraphs:
            joined = preprocess_text("\n\n".join(paragraphs))
            chunks, _, _ = _chunk_page_text(joined, PdfConfig.MIN_CHARS_PER_CHUNK, PdfConfig.MAX_CHARS_PER_CHUNK)
            return chunks
    except Exception:
        pass

    # --- attempt 2: win32com / Microsoft Word COM (Windows + Office) ---
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word_doc = None
        try:
            word_doc = word.Documents.Open(str(local_path.resolve()))
            raw_text = word_doc.Content.Text or ""
        finally:
            if word_doc is not None:
                word_doc.Close(False)
            word.Quit()
        pythoncom.CoUninitialize()

        if raw_text.strip():
            joined = preprocess_text(raw_text)
            chunks, _, _ = _chunk_page_text(joined, PdfConfig.MIN_CHARS_PER_CHUNK, PdfConfig.MAX_CHARS_PER_CHUNK)
            return chunks
    except Exception:
        pass

    raise ValueError(
        "无法解析 .doc 文件：当前环境不支持 Word 97-2003 二进制格式。"
        "请在 Word 中将文件另存为 .docx 格式后重新上传。"
    )


def _parse_pdf_to_chunks(local_path: Path) -> List[str]:
    """Parse PDF text and convert to semantic chunks."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ValueError("pypdf is not installed, cannot parse PDF files")

    reader = PdfReader(str(local_path))
    all_text_parts: List[str] = []
    for page in reader.pages:
        page_text = (page.extract_text() or "").strip()
        if page_text:
            all_text_parts.append(page_text)

    if not all_text_parts:
        return []

    from src.vector_db.pdf2vector import _chunk_page_text, preprocess_text

    joined_text = preprocess_text("\n\n".join(all_text_parts))
    chunks, _, _ = _chunk_page_text(
        joined_text,
        PdfConfig.MIN_CHARS_PER_CHUNK,
        PdfConfig.MAX_CHARS_PER_CHUNK,
    )
    return chunks


def _ingest_excel_to_mysql(local_path: Path) -> int:
    """Parse Excel file and import structured data into MySQL."""
    from src.database.excel2mysql import _read_single_file, _deduplicate_records
    from src.database.mysql_conn import MysqlConnection

    records = _read_single_file(str(local_path))
    if not records:
        return 0

    unique_records = _deduplicate_records(records)
    if not unique_records:
        return 0

    mysql_client = MysqlConnection()
    connection_tuple = mysql_client.connect_db()
    if connection_tuple is None:
        raise RuntimeError("database connection failed")

    connection, cursor = connection_tuple
    insert_sql = (
        "INSERT INTO campus_struct_data (category, item, operation, time_requirement, channel, source_note) "
        "VALUES (%(category)s, %(item)s, %(operation)s, %(time_requirement)s, %(channel)s, %(source_note)s)"
    )
    try:
        cursor.executemany(insert_sql, unique_records)
        connection.commit()
        return cursor.rowcount
    except Exception as import_error:
        connection.rollback()
        raise RuntimeError(f"excel import failed: {import_error}") from import_error
    finally:
        mysql_client.close_db()


def _build_sources(file_name: str, count: int, doc_type: str = "doc") -> List[str]:
    """Build Milvus source labels for chunks."""
    stem = Path(file_name).stem
    return [f"{stem}_{doc_type}_chunk_{index}" for index in range(1, count + 1)]


async def ingest_docx_upload(file: UploadFile) -> dict[str, object]:
    """Handle upload parsing, vectorization, and Milvus insertion (legacy docx endpoint)."""
    return await ingest_file_upload(file)


async def ingest_file_upload(file: UploadFile) -> dict[str, object]:
    """Handle upload of excel/pdf/word: save, parse, and ingest into DB/Milvus."""
    safe_file_name = _validate_file_name(file.filename)
    suffix = Path(safe_file_name).suffix.lower()

    content = await file.read()
    if not content:
        raise ValueError("uploaded file is empty")

    local_path = await asyncio.to_thread(_persist_file_bytes, safe_file_name, content)

    # Persist upload metadata for admin document list page.
    try:
        await asyncio.to_thread(add_pdf_info, safe_file_name, str(local_path))
    except Exception:
        # Keep upload flow available even when metadata storage fails.
        pass

    if suffix in ALLOWED_EXCEL:
        inserted_count = await asyncio.to_thread(_ingest_excel_to_mysql, local_path)
        return {
            "file_name": safe_file_name,
            "chunks": inserted_count,
            "inserted_count": inserted_count,
            "save_path": str(local_path),
        }

    if suffix == ".pdf":
        chunks = await asyncio.to_thread(_parse_pdf_to_chunks, local_path)
        doc_type = "pdf"
    elif suffix == ".doc":
        chunks = await asyncio.to_thread(_parse_doc_to_chunks, local_path)
        doc_type = "doc"
    else:
        chunks = await asyncio.to_thread(_parse_docx_to_chunks, local_path)
        doc_type = "docx"

    if not chunks:
        return {
            "file_name": safe_file_name,
            "chunks": 0,
            "inserted_count": 0,
            "save_path": str(local_path),
        }

    from src.vector_db.milvus_operate import insert_vectors_to_milvus
    from src.vector_db.vector_generator import text_to_vector

    vectors = await asyncio.to_thread(text_to_vector, chunks)
    sources = _build_sources(safe_file_name, len(chunks), doc_type)
    inserted_count = await asyncio.to_thread(insert_vectors_to_milvus, vectors, chunks, sources)

    return {
        "file_name": safe_file_name,
        "chunks": len(chunks),
        "inserted_count": int(inserted_count),
        "save_path": str(local_path),
    }


def list_uploaded_documents() -> list[dict[str, object]]:
    """Return uploaded document records for frontend admin table."""
    try:
        rows = list_pdf_info()
    except Exception:
        return []

    documents: list[dict[str, object]] = []
    for row in rows:
        name = str(row.get("doc_name") or "")
        suffix = Path(name).suffix.replace(".", "").upper() if name else ""
        documents.append(
            {
                "id": int(row.get("id") or 0),
                "file_name": name,
                "file_type": suffix,
                "save_path": str(row.get("doc_path") or ""),
                "uploaded_at": row.get("upload_time"),
            }
        )

    return documents


def delete_uploaded_document(doc_id: int) -> bool:
    """Delete an uploaded document record by id."""
    return del_pdf_info(doc_id)
